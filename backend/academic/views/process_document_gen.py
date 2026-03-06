"""
Generación automática de documentos académicos
═══════════════════════════════════════════════════════════════
IESPP "Gustavo Allende Llavería" — Tarma, Junín

Documentos basados en las plantillas REALES de la institución:
  1. Constancia de Estudios → plantilla .docx + LibreOffice (diseño fiel)
  2. Constancia de Orden de Mérito  → ReportLab
  3. Constancia de Tercio Superior  → ReportLab
  4. Certificado de Egresado (SIA)  → ReportLab
  5. Ficha de Matrícula             → ReportLab

CONSTANCIA DE ESTUDIOS — Flujo:
  La plantilla .docx debe estar en:
    <BASE_DIR>/academic/templates/constancias/constancia_estudios_template.docx
  
  El .docx debe tener estos placeholders (texto exacto):
    {{NOMBRES_APELLIDOS}}   → "Brayan Pedro, MORALES JAJAICUCHO"
    {{DNI}}                 → "78459832"
    {{DOMICILIO}}           → "Flor de Romero Urb. Molino de Vargas"
    {{CIUDAD}}              → "Tarma"
    {{SEMESTRE_RANGE}}      → "VIII"  (se convierte de número a romano)
    {{SECCION}}             → "A"
    {{TURNO}}               → "tarde"
    {{CARRERA}}             → "COMUNICACIÓN"
    {{CODIGO_MATRICULA}}    → "78459832"
    {{DIA}}                 → "11"
    {{MES}}                 → "setiembre"
    {{ANIO}}                → "2025"

  Si la plantilla NO existe → fallback automático al generador ReportLab.
"""

import io
import json
import logging
import os
import shutil
import subprocess
import tempfile
import urllib.request
from datetime import datetime
from collections import OrderedDict

from django.conf import settings
from django.db.models import Avg
from django.shortcuts import get_object_or_404
from django.core.files.base import ContentFile
from rest_framework.views import APIView
from rest_framework import permissions, status
from rest_framework.response import Response
from rest_framework_simplejwt.authentication import JWTAuthentication

from academic.models import AcademicProcess, ProcessFile
from .utils import ok, _can_admin_enroll

logger = logging.getLogger("academic.processes")

# ── python-docx ──
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── Imports seguros ──
try:
    from students.models import Student
except ImportError:
    Student = None

try:
    from academic.models import AcademicGradeRecord
except ImportError:
    AcademicGradeRecord = None

try:
    from academic.models import Plan, PlanCourse, Course
except ImportError:
    Plan = None
    PlanCourse = None
    Course = None

try:
    from catalogs.models import Career
except ImportError:
    Career = None

try:
    from catalogs.models import InstitutionSetting
except ImportError:
    InstitutionSetting = None

# ── ReportLab ──
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image,
    )
    from reportlab.lib import colors
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False


# ═══════════════════════════════════════════════════════════════
# CONSTANTES
# ═══════════════════════════════════════════════════════════════

_DEFAULT_INST = {
    "name":                     "INSTITUTO DE EDUCACIÓN SUPERIOR PEDAGÓGICO PÚBLICO",
    "short_name":               "I.E.S.P.P.",
    "institution_name":         '"GUSTAVO ALLENDE LLAVERÍA"',
    "city":                     "Tarma",
    "region":                   "Junín",
    "province":                 "Tarma",
    "address":                  "Carretera Central Km. 4 S/N Pomachaca",
    "ds_creation":              "D.S. 059-1984-ED",
    "rm_revalidation":          "Reinscripción D.S. 017-2002-ED",
    "modular_code":             "0609370",
    "management":               "Pública",
    "email":                    "iespp.gal.tarma@gmail.com",
    "phone":                    "",
    "director_name":            "",
    "secretary_name":           "",
    "academic_head_name":       "",
    "logo_url":                 "",
    "signature_url":            "",
    "secretary_signature_url":  "",
}

MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}

DOCUMENT_GENERATORS = {}

# ── Ruta a la plantilla .docx de la Constancia de Estudios ──
CONSTANCIA_ESTUDIOS_TEMPLATE = os.path.join(
    settings.BASE_DIR,
    "academic", "templates", "constancias",
    "constancia_estudios_template.docx",
)


def register_generator(kind):
    def decorator(func):
        DOCUMENT_GENERATORS[kind] = func
        return func
    return decorator


# ═══════════════════════════════════════════════════════════════
# HELPERS — INSTITUCIÓN
# ═══════════════════════════════════════════════════════════════

def _get_institution():
    data = dict(_DEFAULT_INST)
    if not InstitutionSetting:
        return data
    try:
        obj = InstitutionSetting.objects.filter(pk=1).first()
        if obj and isinstance(obj.data, dict):
            for k, v in obj.data.items():
                if v:
                    data[k] = v
    except Exception as e:
        logger.warning(f"Error leyendo InstitutionSetting: {e}")
    return data


def _resolve_media_path(url_or_path):
    if not url_or_path:
        return None
    p = str(url_or_path)
    if p.startswith(("http://", "https://")):
        try:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            urllib.request.urlretrieve(p, tmp.name)
            return tmp.name
        except Exception as e:
            logger.warning(f"No se pudo descargar media {p}: {e}")
            return None
    if "/media/" in p:
        rel = p.split("/media/")[-1]
        full = os.path.join(settings.MEDIA_ROOT, rel)
        if os.path.exists(full):
            return full
    if p.startswith("/") and os.path.exists(p):
        return p
    full = os.path.join(settings.MEDIA_ROOT, p.lstrip("/"))
    if os.path.exists(full):
        return full
    return None


def _load_image(inst, key, max_w=3 * cm, max_h=2 * cm):
    if not HAS_REPORTLAB:
        return None
    path = _resolve_media_path(inst.get(key))
    if not path:
        return None
    try:
        img = Image(path, width=max_w, height=max_h)
        img.hAlign = "CENTER"
        return img
    except Exception as e:
        logger.warning(f"Error cargando imagen ({key}): {e}")
        return None


# ═══════════════════════════════════════════════════════════════
# HELPERS — ESTUDIANTE
# ═══════════════════════════════════════════════════════════════

def _get_student(student_id):
    data = {
        "id": student_id,
        "nombre_completo": f"Estudiante {student_id}",
        "apellidos": "",
        "nombres": "",
        "dni": "",
        "codigo": "",
        "sexo": "",
        "email": "",
        "carrera": "",
        "carrera_id": None,
        "plan_name": "",
        "plan_id": None,
        "ciclo": "",
        "seccion": "",
        "turno": "",
        "periodo": "",
        "domicilio": "",
        "promocion": "",
    }
    if not Student:
        return data
    try:
        st = Student.objects.select_related("plan", "plan__career", "user").get(id=student_id)

        data["nombres"] = getattr(st, "nombres", "") or ""
        ap_pat = getattr(st, "apellido_paterno", "") or ""
        ap_mat = getattr(st, "apellido_materno", "") or ""
        data["apellidos"] = f"{ap_pat} {ap_mat}".strip()

        if data["apellidos"] and data["nombres"]:
            data["nombre_completo"] = f'{data["apellidos"].upper()}, {data["nombres"]}'
        elif data["nombres"]:
            data["nombre_completo"] = data["nombres"]

        data["dni"]    = getattr(st, "num_documento", "") or ""
        data["codigo"] = getattr(st, "codigo_modular", "") or data["dni"]
        data["sexo"]   = getattr(st, "sexo", "") or ""
        data["email"]  = getattr(st, "email", "") or ""

        if hasattr(st, "plan") and st.plan:
            plan = st.plan
            data["plan_name"] = plan.name or ""
            data["plan_id"]   = plan.id
            if hasattr(plan, "career") and plan.career:
                data["carrera"]    = plan.career.name or ""
                data["carrera_id"] = plan.career.id
        if not data["carrera"]:
            data["carrera"] = getattr(st, "programa_carrera", "") or ""

        data["ciclo"]   = str(getattr(st, "ciclo", "") or "")
        data["seccion"] = getattr(st, "seccion", "") or ""
        data["turno"]   = getattr(st, "turno", "") or ""
        data["periodo"] = getattr(st, "periodo", "") or ""

        parts = []
        for f in ("distrito", "provincia", "region"):
            v = getattr(st, f, "") or ""
            if v:
                parts.append(v)
        data["domicilio"] = ", ".join(parts) if parts else ""

        data["promocion"] = (
            getattr(st, "promocion", "") or
            getattr(st, "promotion_year", "") or ""
        )

    except Student.DoesNotExist:
        logger.warning(f"Estudiante {student_id} no encontrado")
    except Exception as e:
        logger.warning(f"Error datos estudiante {student_id}: {e}")
    return data


def _get_grades_by_term(student_id):
    from collections import OrderedDict
    grades = OrderedDict()
    if not AcademicGradeRecord:
        return grades
    try:
        results = (
            AcademicGradeRecord.objects
            .filter(student_id=student_id)
            .values("term")
            .annotate(promedio=Avg("final_grade"))
            .order_by("term")
        )
        for r in results:
            term = str(r["term"] or "")
            if term:
                grades[term] = round(r["promedio"] or 0, 2)
    except Exception as e:
        logger.warning(f"Error grades estudiante {student_id}: {e}")
    return grades


def _get_enrolled_courses(plan_id, semester_filter=None):
    courses = []
    if not PlanCourse or not plan_id:
        return courses
    try:
        qs = (
            PlanCourse.objects
            .filter(plan_id=plan_id)
            .select_related("course")
            .order_by("semester", "display_code", "id")
        )
        if semester_filter:
            try:
                qs = qs.filter(semester=int(semester_filter))
            except (ValueError, TypeError):
                pass
        for pc in qs:
            course = pc.course
            courses.append({
                "nombre":   pc.display_name or (course.name if course else "—"),
                "codigo":   pc.display_code or (course.code if course else ""),
                "horas":    pc.weekly_hours or 0,
                "creditos": pc.credits or (course.credits if course else 0),
                "ciclo":    pc.semester or 0,
                "tipo":     pc.type or "MANDATORY",
            })
    except Exception as e:
        logger.warning(f"Error courses plan {plan_id}: {e}")
    return courses


def _get_extra_data(proc):
    if not hasattr(proc, "extra_data") or not proc.extra_data:
        return {}
    try:
        return (
            json.loads(proc.extra_data)
            if isinstance(proc.extra_data, str)
            else (proc.extra_data or {})
        )
    except Exception:
        return {}


# ═══════════════════════════════════════════════════════════════
# HELPER — NÚMERO A ROMANO
# ═══════════════════════════════════════════════════════════════

def _to_roman(n) -> str:
    try:
        n = int(n)
    except (ValueError, TypeError):
        return str(n)
    vals = [
        (10, "X"), (9, "IX"), (8, "VIII"), (7, "VII"), (6, "VI"),
        (5, "V"),  (4, "IV"), (3, "III"),  (2, "II"),  (1, "I"),
    ]
    result = ""
    for v, s in vals:
        while n >= v:
            result += s
            n -= v
    return result or str(n)


# ═══════════════════════════════════════════════════════════════
# GENERADOR DOCX — CONSTANCIA DE ESTUDIOS
# Usa la plantilla .docx real + LibreOffice → PDF
# ═══════════════════════════════════════════════════════════════

def _replace_in_paragraph(paragraph, replacements: dict):
    """
    Reemplaza placeholders en un párrafo manejando el caso donde
    el placeholder está partido entre varios runs del XML.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(key in full_text for key in replacements):
        return
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, str(value))
    # Poner todo el texto en el primer run, vaciar el resto
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_docx(doc, replacements: dict):
    """Recorre todo el documento (cuerpo + tablas + header/footer) y aplica reemplazos."""
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)

    return doc


def _find_libreoffice() -> str | None:
    candidates = [
        "libreoffice", "soffice",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    for cmd in candidates:
        if shutil.which(cmd):
            return cmd
    return None


def _docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convierte bytes de un .docx a bytes de PDF usando LibreOffice."""
    lo = _find_libreoffice()
    if not lo:
        raise RuntimeError(
            "LibreOffice no está instalado. "
            "Instala con: sudo apt install libreoffice"
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "constancia.docx")
        pdf_path  = os.path.join(tmpdir, "constancia.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [lo, "--headless", "--norestore",
             "--convert-to", "pdf",
             "--outdir", tmpdir, docx_path],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice falló: {result.stderr or result.stdout}"
            )

        if not os.path.exists(pdf_path):
            raise RuntimeError(
                f"LibreOffice no generó el PDF. stdout: {result.stdout}"
            )

        with open(pdf_path, "rb") as f:
            return f.read()


def _generate_constancia_docx(process, student: dict, extra: dict, inst: dict) -> tuple:
    """
    Genera la Constancia de Estudios usando la plantilla .docx institucional.
    Retorna (BytesIO del PDF, filename) o lanza excepción.
    """
    if not HAS_DOCX:
        raise ImportError("Instala python-docx: pip install python-docx")

    now = datetime.now()

    # ── Preparar datos ──
    nombres   = student.get("nombres", "")
    apellidos = student.get("apellidos", "")
    sexo      = student.get("sexo", "").upper()
    ciclo     = student.get("ciclo", "") or extra.get("semester", "")
    seccion   = student.get("seccion", "") or extra.get("section", "")
    turno     = student.get("turno", "") or extra.get("shift", "")
    carrera   = student.get("carrera", "") or extra.get("career", "")
    dni       = student.get("dni", "")
    codigo    = student.get("codigo", "") or dni
    domicilio = student.get("domicilio", "") or extra.get("address", "")
    city      = inst.get("city", "Tarma")

    # Nombre: "Nombres, APELLIDOS"
    if nombres and apellidos:
        nombre_display = f"{nombres}, {apellidos.upper()}"
    else:
        nombre_display = student.get("nombre_completo", "")

    # Semestre a romano
    ciclo_romano = _to_roman(ciclo) if str(ciclo).isdigit() else str(ciclo)

    # ── Diccionario de reemplazos ──
    replacements = {
        "{{NOMBRES_APELLIDOS}}": nombre_display,
        "{{DNI}}":               dni,
        "{{DOMICILIO}}":         domicilio,
        "{{CIUDAD}}":            city,
        "{{SEMESTRE_RANGE}}":    ciclo_romano,
        "{{SECCION}}":           seccion,
        "{{TURNO}}":             turno,
        "{{CARRERA}}":           carrera.upper() if carrera else "",
        "{{CODIGO_MATRICULA}}":  codigo,
        "{{DIA}}":               str(now.day),
        "{{MES}}":               MESES.get(now.month, ""),
        "{{ANIO}}":              str(now.year),
        # aliases por si el template usa otros nombres
        "{{NOMBRE}}":            nombre_display,
        "{{SEMESTRE}}":          ciclo_romano,
        "{{CODIGO}}":            codigo,
    }

    # ── Abrir plantilla, reemplazar, guardar en memoria ──
    doc = DocxDocument(CONSTANCIA_ESTUDIOS_TEMPLATE)
    doc = _replace_in_docx(doc, replacements)

    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    # ── Convertir a PDF con LibreOffice ──
    pdf_bytes = _docx_bytes_to_pdf_bytes(docx_bytes)

    filename = f"CONSTANCIA-ESTUDIOS_{process.id:05d}_{now.strftime('%Y%m%d')}.pdf"
    buf = io.BytesIO(pdf_bytes)
    buf.seek(0)
    return buf, filename


# ═══════════════════════════════════════════════════════════════
# ESTILOS REPORTLAB (para los demás documentos)
# ═══════════════════════════════════════════════════════════════

def _get_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "InstName", parent=styles["Heading1"],
        fontSize=12, alignment=TA_CENTER, spaceAfter=0, spaceBefore=0,
        textColor=colors.HexColor("#1a237e"), fontName="Helvetica-Bold", leading=15,
    ))
    styles.add(ParagraphStyle(
        "InstSub", parent=styles["Normal"],
        fontSize=9, alignment=TA_CENTER, spaceAfter=2,
        textColor=colors.HexColor("#37474f"),
    ))
    styles.add(ParagraphStyle(
        "DocTitle", parent=styles["Heading2"],
        fontSize=15, alignment=TA_CENTER, spaceBefore=14, spaceAfter=14,
        textColor=colors.black, fontName="Helvetica-Bold", leading=20,
    ))
    styles.add(ParagraphStyle(
        "BodyJ", parent=styles["Normal"],
        fontSize=11, alignment=TA_JUSTIFY, leading=16, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "BodyC", parent=styles["Normal"],
        fontSize=11, alignment=TA_CENTER, leading=16, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "SmallC", parent=styles["Normal"],
        fontSize=8, alignment=TA_CENTER,
        textColor=colors.HexColor("#616161"),
    ))
    styles.add(ParagraphStyle(
        "Right", parent=styles["Normal"],
        fontSize=11, alignment=TA_RIGHT,
    ))
    styles.add(ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontSize=7, alignment=TA_CENTER,
        textColor=colors.HexColor("#9e9e9e"),
    ))
    styles.add(ParagraphStyle(
        "FirmaLabel", parent=styles["Normal"],
        fontSize=9, alignment=TA_CENTER,
        textColor=colors.HexColor("#333333"),
    ))
    return styles


# ═══════════════════════════════════════════════════════════════
# BLOQUES REUTILIZABLES REPORTLAB
# ═══════════════════════════════════════════════════════════════

def _header(styles, inst):
    elems = []
    logo = _load_image(inst, "logo_url", 3.5 * cm, 2.5 * cm)
    name    = inst.get("name", _DEFAULT_INST["name"])
    nombre  = inst.get("institution_name", _DEFAULT_INST["institution_name"])
    city    = inst.get("city", "Tarma")
    region  = inst.get("region", "Junín")
    ds      = inst.get("ds_creation", _DEFAULT_INST["ds_creation"])
    rm      = inst.get("rm_revalidation", "")
    address = inst.get("address", "")
    email   = inst.get("email", "")

    text = [
        Paragraph(name, styles["InstName"]),
        Paragraph(nombre, styles["InstName"]),
        Paragraph(f"{city} - {region} - PERÚ", styles["InstSub"]),
        Paragraph(f"{ds}{' - ' + rm if rm else ''}", styles["SmallC"]),
    ]
    if address:
        text.append(Paragraph(f"{address}{' - Email: ' + email if email else ''}", styles["SmallC"]))

    if logo:
        ht = Table([[logo, text]], colWidths=[4 * cm, 12 * cm])
        ht.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN",  (0, 0), (0, 0), "CENTER"),
            ("ALIGN",  (1, 0), (1, 0), "CENTER"),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        elems.append(ht)
    else:
        elems.extend(text)

    elems.append(Spacer(1, 6))
    line = Table([[""]], colWidths=[16 * cm], rowHeights=[2])
    line.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#1565c0"))]))
    elems.append(line)
    elems.append(Spacer(1, 10))
    return elems


def _firma(styles, inst, cargo="Director General", sig_key="signature_url"):
    elems = [Spacer(1, 30)]
    short  = inst.get("short_name", _DEFAULT_INST["short_name"])
    nombre = inst.get("institution_name", _DEFAULT_INST["institution_name"])

    firmante = ""
    if "director" in cargo.lower():
        firmante = inst.get("director_name", "")
    elif "secretari" in cargo.lower():
        firmante = inst.get("secretary_name", "")
    elif "unidad" in cargo.lower() or "jefe" in cargo.lower():
        firmante = inst.get("academic_head_name", "")

    img  = _load_image(inst, sig_key, 3 * cm, 1.5 * cm)
    rows = []
    if img:
        rows.append([img])
    rows.append(["_" * 35])
    if firmante:
        rows.append([Paragraph(f"<b>{firmante}</b>", styles["FirmaLabel"])])
    rows.append([Paragraph(cargo.upper(), styles["FirmaLabel"])])
    rows.append([Paragraph(f"{short} {nombre}", styles["FirmaLabel"])])

    t = Table(rows, colWidths=[8 * cm])
    t.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING",    (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
    ]))
    outer = Table([[t]], colWidths=[16 * cm])
    outer.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER")]))
    elems.append(outer)
    return elems


def _firma_triple(styles, inst):
    elems   = [Spacer(1, 20)]
    dir_img = _load_image(inst, "signature_url", 2.5 * cm, 1.2 * cm)
    sec_img = _load_image(inst, "secretary_signature_url", 2.5 * cm, 1.2 * cm)

    def _col(img, name, cargo):
        rows = []
        if img:
            rows.append([img])
        rows.append(["_" * 25])
        if name:
            rows.append([Paragraph(f"<b>{name}</b>", styles["SmallC"])])
        rows.append([Paragraph(cargo, styles["SmallC"])])
        t = Table(rows, colWidths=[5 * cm])
        t.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING",    (0, 0), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ]))
        return t

    f1 = _col(dir_img, inst.get("director_name", ""),     "DIRECTORA GENERAL")
    f2 = _col(sec_img, inst.get("secretary_name", ""),    "SECRETARIO ACADÉMICO")
    f3 = _col(None,    inst.get("academic_head_name", ""), "JEFE UNIDAD ACADÉMICA")

    outer = Table([[f1, f2, f3]], colWidths=[5.5 * cm, 5.5 * cm, 5.5 * cm])
    outer.setStyle(TableStyle([
        ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
    ]))
    elems.append(outer)
    return elems


def _footer(styles, process_id):
    now = datetime.now()
    return [
        Spacer(1, 16),
        Paragraph(
            f"Documento generado por el Sistema Académico | "
            f"Proceso #{process_id:05d} | {now.strftime('%d/%m/%Y %H:%M')}",
            styles["Footer"],
        ),
    ]


def _fecha(inst=None):
    now  = datetime.now()
    city = (inst or {}).get("city", "Tarma")
    return f"{city}, {now.day} de {MESES.get(now.month, '')} del {now.year}"


# ═══════════════════════════════════════════════════════════════
#  1. CONSTANCIA DE ESTUDIOS
#  Fallback ReportLab (cuando no existe la plantilla .docx)
# ═══════════════════════════════════════════════════════════════

@register_generator("CONSTANCIA_ESTUDIOS")
def _gen_constancia_estudios(process, student, extra, styles, inst):
    elems = _header(styles, inst)
    elems.append(Paragraph("CONSTANCIA  DE  ESTUDIOS", styles["DocTitle"]))
    elems.append(Spacer(1, 4))

    i_name   = inst.get("name", _DEFAULT_INST["name"])
    i_nombre = inst.get("institution_name", _DEFAULT_INST["institution_name"])
    city     = inst.get("city", "Tarma")

    elems.append(Paragraph(
        f'QUIEN SUSCRIBE LA SECRETARIA ACADÉMICA DEL {i_name} '
        f'{i_nombre} de {city}:',
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 4))
    elems.append(Paragraph("<b>HACE  CONSTAR</b>", styles["BodyC"]))
    elems.append(Spacer(1, 8))

    ciclo     = student["ciclo"] or extra.get("semester", "X")
    seccion   = student["seccion"] or extra.get("section", "")
    turno     = student["turno"] or extra.get("shift", "")
    domicilio = student["domicilio"] or extra.get("address", "")
    carrera   = student["carrera"] or extra.get("career", "—")
    codigo    = student["codigo"] or extra.get("code", "")

    ciclo_romano = _to_roman(ciclo) if str(ciclo).isdigit() else str(ciclo)

    nombres   = student.get("nombres", "")
    apellidos = student.get("apellidos", "")
    nombre_display = f"{nombres}, {apellidos.upper()}" if apellidos else student["nombre_completo"]

    body = (
        f'Que, <b>{nombre_display}</b> es estudiante regular de esta '
        f'Casa de Formación Magisterial identificado(a) con '
        f'D.N.I Nº <b>{student["dni"]}</b>'
    )
    if domicilio:
        body += f' con domicilio en {domicilio} de la ciudad de {city},'
    body += (
        f' quien cursa estudios superiores a partir del I al '
        f'<b>{ciclo_romano}</b> Semestre Académico'
    )
    if seccion:
        body += f' Sección "{seccion}"'
    if turno:
        body += f' Turno {turno}'
    body += f' en el Programa de Estudios de <b>{carrera.upper()}</b>'
    if codigo:
        body += f' con código de matrícula Nº <b>{codigo}</b>'
    body += "."

    elems.append(Paragraph(body, styles["BodyJ"]))
    elems.append(Spacer(1, 16))
    elems.append(Paragraph(
        "Se expide la presente a solicitud del interesado(a), para los fines "
        "que estime por conveniente.",
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 8))
    elems.append(Paragraph(_fecha(inst), styles["Right"]))
    elems += _firma(styles, inst, "Secretaria Académica", "secretary_signature_url")
    elems += _footer(styles, process.id)
    return elems


# ═══════════════════════════════════════════════════════════════
#  2. CONSTANCIA DE ORDEN DE MÉRITO
# ═══════════════════════════════════════════════════════════════

@register_generator("CONSTANCIA_ORDEN_MERITO")
def _gen_constancia_orden_merito(process, student, extra, styles, inst):
    elems = _header(styles, inst)
    elems.append(Paragraph("Constancia  de  Orden de Mérito", styles["DocTitle"]))

    i_name   = inst.get("name", _DEFAULT_INST["name"])
    i_nombre = inst.get("institution_name", _DEFAULT_INST["institution_name"])
    city     = inst.get("city", "Tarma")

    elems.append(Paragraph(
        f'QUIEN SUSCRIBE EL SECRETARIO ACADÉMICO DEL {i_name} '
        f'{i_nombre} - {city.upper()}:',
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 4))
    elems.append(Paragraph("<b>HACE CONSTAR QUE:</b>", styles["BodyC"]))
    elems.append(Spacer(1, 8))

    puesto          = extra.get("rank", extra.get("puesto", "PRIMER"))
    carrera         = student["carrera"] or extra.get("career", "—")
    semester_range  = extra.get("semester_range", "I al X")
    promocion       = extra.get("promotion_year", extra.get("promocion", ""))

    body = (
        f'Que, <b>{student["nombre_completo"]}</b> identificado(a) con D.N.I Nº '
        f'<b>{student["dni"]}</b>, ha ocupado el <b>{puesto} LUGAR</b> en el orden de '
        f'mérito de la especialidad de <b>{carrera.upper()}</b> durante sus estudios de '
        f'Educación Superior del <b>{semester_range}</b> Semestre Académico '
        f'obteniendo los siguientes Promedios Ponderado:'
    )
    elems.append(Paragraph(body, styles["BodyJ"]))
    elems.append(Spacer(1, 8))

    grades = _get_grades_by_term(process.student_id)
    if not grades and extra.get("grades"):
        try:
            for k, v in extra["grades"].items():
                grades[k] = float(v)
        except Exception:
            pass

    if grades:
        items = list(grades.items())
        mid   = (len(items) + 1) // 2
        left  = items[:mid]
        right = items[mid:]

        rows = [["Semestre", "Prom. Pond.", "", "Semestre", "Prom. Pond."]]
        for i in range(max(len(left), len(right))):
            l_s = left[i][0]  if i < len(left)  else ""
            l_v = f"{left[i][1]:.2f}"  if i < len(left)  else ""
            r_s = right[i][0] if i < len(right) else ""
            r_v = f"{right[i][1]:.2f}" if i < len(right) else ""
            rows.append([l_s, l_v, "", r_s, r_v])

        gt = Table(rows, colWidths=[3 * cm, 2.5 * cm, 0.5 * cm, 3 * cm, 2.5 * cm])
        gt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e3f2fd")),
            ("FONTSIZE",   (0, 0), (-1, -1), 10),
            ("GRID",       (0, 0), (1, -1), 0.5, colors.HexColor("#bdbdbd")),
            ("GRID",       (3, 0), (4, -1), 0.5, colors.HexColor("#bdbdbd")),
            ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elems.append(gt)
        elems.append(Spacer(1, 10))

    prom_sem = extra.get("semestral_average", "")
    if prom_sem:
        txt = f'Promedio Ponderado semestral de <b>{prom_sem}</b>'
        if promocion:
            txt += f' en la promoción {promocion}'
        elems.append(Paragraph(txt, styles["BodyC"]))
        elems.append(Spacer(1, 6))

    elems.append(Paragraph(
        'Según consta en las Actas Consolidadas de Evaluación de esta '
        'Institución, a los que me remito en caso de ser necesario.',
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 6))
    elems.append(Paragraph(
        "Se expide la presente a solicitud del interesado(a), para los fines "
        "que estime conveniente.",
        styles["BodyJ"],
    ))
    elems.append(Paragraph(_fecha(inst), styles["Right"]))
    elems += _firma(styles, inst, "Secretario(a) Académico(a)", "secretary_signature_url")
    elems += _footer(styles, process.id)
    return elems


# ═══════════════════════════════════════════════════════════════
#  3. CONSTANCIA DE TERCIO SUPERIOR
# ═══════════════════════════════════════════════════════════════

@register_generator("CONSTANCIA_TERCIO")
def _gen_constancia_tercio(process, student, extra, styles, inst):
    elems = _header(styles, inst)
    elems.append(Paragraph("CONSTANCIA  DE  TERCIO SUPERIOR", styles["DocTitle"]))

    i_name   = inst.get("name", _DEFAULT_INST["name"])
    i_nombre = inst.get("institution_name", _DEFAULT_INST["institution_name"])
    city     = inst.get("city", "Tarma")

    elems.append(Paragraph(
        f'LA DIRECTORA GENERAL DEL {i_name} {i_nombre} de {city}:',
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 8))
    elems.append(Paragraph('Certifica que el(a) egresada:', styles["BodyJ"]))
    elems.append(Spacer(1, 4))
    elems.append(Paragraph(f'<b>{student["nombre_completo"]}</b>', styles["BodyC"]))
    elems.append(Spacer(1, 8))
    elems.append(Paragraph(
        'Pertenece al Tercio Superior de su promoción por los promedios obtenidos '
        'al término de los cinco años de estudios profesionales.',
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 8))

    carrera       = student["carrera"] or extra.get("career", "—")
    anio          = extra.get("academic_year", extra.get("years", ""))
    num_egresados = extra.get("total_graduates", extra.get("num_egresados", ""))
    puesto        = extra.get("rank", extra.get("puesto", ""))
    prom_tercio   = extra.get("tercio_average", extra.get("promedio_tercio", ""))
    promocion     = extra.get("promotion_year", extra.get("promocion", ""))

    info = [
        ["Carrera Profesional",                          str(carrera)],
        ["Años Académicos",                              str(anio) if anio else ""],
        ["Nº de egresados del cuadro de mérito promocional", str(num_egresados) if num_egresados else ""],
        ["Puesto del egresado en su promoción",          str(puesto) if puesto else ""],
        ["Promedio del tercio superior",                 str(prom_tercio) if prom_tercio else ""],
    ]
    it = Table(info, colWidths=[9 * cm, 6 * cm])
    it.setStyle(TableStyle([
        ("FONTSIZE",  (0, 0), (-1, -1), 10),
        ("GRID",      (0, 0), (-1, -1), 0.5, colors.HexColor("#bdbdbd")),
        ("BACKGROUND",(0, 0), (0, -1),  colors.HexColor("#e3f2fd")),
        ("FONTNAME",  (0, 0), (0, -1),  "Helvetica-Bold"),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
    ]))
    elems.append(it)
    elems.append(Spacer(1, 12))
    elems.append(Paragraph(
        f'Según consta en el Archivo de Orden de Mérito de la promoción '
        f'{promocion or "____"} de esta Institución, a los que me remito '
        f'en caso de ser necesario.',
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 6))
    elems.append(Paragraph(
        "Se expide la presente a solicitud del interesado(a), para los fines "
        "que estime conveniente.",
        styles["BodyJ"],
    ))
    elems.append(Paragraph(_fecha(inst), styles["Right"]))
    elems += _firma(styles, inst, "Directora General", "signature_url")
    elems += _footer(styles, process.id)
    return elems


# ═══════════════════════════════════════════════════════════════
#  4. CERTIFICADO DE EGRESADO (SIA)
# ═══════════════════════════════════════════════════════════════

@register_generator("CERTIFICADO_EGRESADO")
def _gen_certificado_egresado(process, student, extra, styles, inst):
    elems = []
    elems.append(Paragraph("MINISTERIO DE EDUCACIÓN", styles["InstSub"]))
    elems.append(Paragraph("DIRECCIÓN DE FORMACIÓN INICIAL DOCENTE", styles["SmallC"]))
    elems.append(Spacer(1, 6))
    elems.extend(_header(styles, inst))

    doc_num = f"N° {process.id:06d}"
    elems.append(Paragraph(f"<b>{doc_num}</b>", styles["Right"]))
    elems.append(Spacer(1, 6))

    i_name   = inst.get("name", _DEFAULT_INST["name"])
    i_nombre = inst.get("institution_name", _DEFAULT_INST["institution_name"])
    city     = inst.get("city", "Tarma")

    elems.append(Paragraph(
        f'El Director General del {i_name} {i_nombre} de {city.upper()} '
        f'que al final suscribe, otorga el presente:',
        styles["BodyJ"],
    ))
    elems.append(Spacer(1, 8))
    elems.append(Paragraph("CERTIFICADO  EGRESADO", styles["DocTitle"]))
    elems.append(Spacer(1, 6))
    elems.append(Paragraph(
        f'<b>A {student["nombre_completo"].upper()}</b>',
        ParagraphStyle(
            "EgName", parent=styles["BodyC"],
            fontSize=13, fontName="Helvetica-Bold",
            spaceBefore=4, spaceAfter=10,
        ),
    ))

    carrera       = student["carrera"] or extra.get("career", "—")
    creditos      = extra.get("credits", extra.get("creditos", "220"))
    promocion     = extra.get("promotion_year", extra.get("promocion", ""))
    periodo_promo = extra.get("promotion_period", "")

    body = (
        f'Por haber concluido y aprobado <b>{creditos}</b> créditos del '
        f'Plan de Estudios de su Formación Inicial Docente, en la carrera de: '
        f'<b>{carrera.upper()}</b>'
    )
    if promocion:
        body += f', Promoción <b>{promocion}'
        if periodo_promo:
            body += f'-{periodo_promo}'
        body += '</b>'
    body += '.'

    elems.append(Paragraph(body, styles["BodyJ"]))
    elems.append(Spacer(1, 16))
    elems.append(Paragraph(f'Dado en {_fecha(inst)}', styles["BodyC"]))
    elems.append(Spacer(1, 10))
    elems += _firma_triple(styles, inst)
    elems += _footer(styles, process.id)
    return elems


# ═══════════════════════════════════════════════════════════════
#  5. FICHA DE MATRÍCULA
# ═══════════════════════════════════════════════════════════════

@register_generator("FICHA_MATRICULA")
def _gen_ficha_matricula(process, student, extra, styles, inst):
    elems  = []
    period = extra.get("period", student["periodo"] or "2025-I")
    ciclo  = student["ciclo"] or extra.get("cycle", "I")
    seccion= student["seccion"] or extra.get("section", "A")
    ciclo_seccion = extra.get("cycle_section", f'{ciclo} - "{seccion}"')

    short      = inst.get("short_name", _DEFAULT_INST["short_name"])
    nombre_inst= inst.get("institution_name", _DEFAULT_INST["institution_name"])
    region     = inst.get("region", "Junín")
    modular    = inst.get("modular_code", _DEFAULT_INST["modular_code"])
    gestion    = inst.get("management", _DEFAULT_INST["management"])
    address    = inst.get("address", "")
    city       = inst.get("city", "Tarma")

    elems.append(Paragraph(
        "FICHA DE MATRÍCULA",
        ParagraphStyle("FT", parent=styles["DocTitle"], fontSize=16, spaceBefore=4, spaceAfter=8),
    ))

    b1 = [
        [Paragraph("<b>Institución</b>", styles["Normal"]),
         Paragraph(f"{short} {nombre_inst}", styles["Normal"]),
         Paragraph("<b>DREJ</b>", styles["Normal"]),
         Paragraph(region, styles["Normal"])],
        [Paragraph("<b>Código Modular</b>", styles["Normal"]),
         Paragraph(modular, styles["Normal"]),
         Paragraph("<b>Gestión</b>", styles["Normal"]),
         Paragraph(gestion, styles["Normal"])],
        [Paragraph("<b>Dirección</b>", styles["Normal"]),
         Paragraph(address, styles["Normal"]),
         Paragraph("<b>Provincia</b>", styles["Normal"]),
         Paragraph(f"{city} / {region}", styles["Normal"])],
    ]
    t1 = Table(b1, colWidths=[3.5 * cm, 5.5 * cm, 3 * cm, 4 * cm])
    t1.setStyle(TableStyle([
        ("FONTSIZE",   (0, 0), (-1, -1), 8),
        ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#bdbdbd")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8eaf6")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#e8eaf6")),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING",   (0, 0), (-1, -1), 4),
    ]))
    elems.append(t1)
    elems.append(Spacer(1, 8))

    carrera   = student["carrera"] or extra.get("career", "—")
    codigo    = student["codigo"] or extra.get("code", "")
    resolucion= extra.get("resolucion", "R.V. Nº 143-2020-MINEDU")
    nombres   = student.get("nombres", "")
    apellidos = student.get("apellidos", "")
    nombre_ficha = f"{apellidos.upper()} {nombres}" if apellidos else student["nombre_completo"]

    b2 = [
        [Paragraph("<b>Programa de Estudios</b>", styles["Normal"]),
         Paragraph(f"<b>{carrera}</b>", styles["Normal"]),
         Paragraph("<b>Período Académico</b>", styles["Normal"]),
         Paragraph(f"<b>{period}</b>", styles["Normal"])],
        [Paragraph("<b>Resolución Autorización</b>", styles["Normal"]),
         Paragraph(resolucion, styles["Normal"]),
         Paragraph("<b>Ciclo - Sección</b>", styles["Normal"]),
         Paragraph(ciclo_seccion, styles["Normal"])],
        [Paragraph("<b>Nombres y Apellidos</b>", styles["Normal"]),
         Paragraph(f"<b>{nombre_ficha}</b>", styles["Normal"]),
         Paragraph("<b>CÓDIGO</b>", styles["Normal"]),
         Paragraph(f"<b>{codigo}</b>", styles["Normal"])],
    ]
    t2 = Table(b2, colWidths=[3.5 * cm, 5.5 * cm, 3.5 * cm, 3.5 * cm])
    t2.setStyle(TableStyle([
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#bdbdbd")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8eaf6")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#e8eaf6")),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING",   (0, 0), (-1, -1), 4),
    ]))
    elems.append(t2)
    elems.append(Spacer(1, 10))

    plan_id  = student.get("plan_id")
    ciclo_int = None
    try:
        ciclo_int = int(ciclo)
    except (ValueError, TypeError):
        pass

    courses = _get_enrolled_courses(plan_id, ciclo_int)
    if not courses:
        for c in (extra.get("courses_list", []) or []):
            if isinstance(c, dict):
                courses.append(c)

    rows    = [["N°", "ASIGNATURA", "HORAS", "CRÉDITOS"]]
    total_h = 0
    total_c = 0
    for i, c in enumerate(courses, 1):
        h  = c.get("horas", 0) or 0
        cr = c.get("creditos", 0) or 0
        total_h += h
        total_c += cr
        rows.append([str(i), c.get("nombre", "—"), str(h), str(cr)])
    if not courses:
        for i in range(1, 9):
            rows.append([str(i), "", "", ""])
    rows.append(["", "TOTAL", str(total_h) if total_h else "", str(total_c) if total_c else ""])

    ct = Table(rows, colWidths=[1 * cm, 10 * cm, 2.5 * cm, 2.5 * cm])
    ct.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1565c0")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#bdbdbd")),
        ("ALIGN",      (0, 0), (0, -1), "CENTER"),
        ("ALIGN",      (2, 0), (3, -1), "CENTER"),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#e3f2fd")),
        ("FONTNAME",   (0, -1), (-1, -1), "Helvetica-Bold"),
    ]))
    elems.append(ct)
    elems.append(Spacer(1, 8))

    sub_rows = [["N°", "CURSOS DE SUBSANACIÓN", "HORAS", "CRÉDITOS"]]
    subs = extra.get("subsanacion_courses", [])
    if subs:
        for i, c in enumerate(subs, 1):
            sub_rows.append([str(i), c.get("nombre", ""),
                             str(c.get("horas", "")), str(c.get("creditos", ""))])
    else:
        sub_rows.append(["1", "", "", ""])

    st = Table(sub_rows, colWidths=[1 * cm, 10 * cm, 2.5 * cm, 2.5 * cm])
    st.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#ef6c00")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#bdbdbd")),
        ("ALIGN",      (0, 0), (0, -1), "CENTER"),
        ("ALIGN",      (2, 0), (3, -1), "CENTER"),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elems.append(st)
    elems += _footer(styles, process.id)
    return elems


# ═══════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL: GENERAR PDF
# ═══════════════════════════════════════════════════════════════

def generate_process_document(process, document_type=None):
    """
    Genera el PDF para un proceso académico.

    CONSTANCIA_ESTUDIOS:
      → plantilla .docx institucional + LibreOffice (diseño fiel al doc original)
      → fallback ReportLab si no existe la plantilla o falta LibreOffice

    FICHA_MATRICULA:
      → HTML + WeasyPrint (diseño tabular fiel al XLS original, A4 horizontal)
      → fallback ReportLab si WeasyPrint no está instalado

    Resto de documentos → ReportLab directamente.

    Retorna (BytesIO, filename).
    """
    from .ficha_matricula_generator import (
        generate_ficha_matricula_weasyprint,
        HAS_WEASYPRINT as _WEASYPRINT_OK,
        _ciclo_int,
    )
    from .orden_merito_generator import (
        generate_orden_merito_weasyprint,
        HAS_WEASYPRINT as _WP_ORDEN,
    )
    from .tercio_superior_generator import (
        generate_tercio_superior_weasyprint,
        HAS_WEASYPRINT as _WP_TERCIO,
    )
    from .certificado_egresado_generator import (
        generate_certificado_egresado_weasyprint,
        HAS_WEASYPRINT as _WP_CERT,
    )

    doc_type = (document_type or process.kind or "").upper().strip()

    # ── CERTIFICADO_EGRESADO → HTML + WeasyPrint (landscape, fiel al SIA) ──
    if doc_type == "CERTIFICADO_EGRESADO":
        if _WP_CERT:
            logger.info(f"Proceso {process.id}: generando CERTIFICADO_EGRESADO con WeasyPrint")
            inst    = _get_institution()
            student = _get_student(process.student_id)
            extra   = _get_extra_data(process)
            return generate_certificado_egresado_weasyprint(process, student, extra, inst)
        else:
            logger.warning("WeasyPrint no instalado → usando ReportLab para CERTIFICADO_EGRESADO")

    # ── CONSTANCIA_TERCIO → HTML + WeasyPrint (cálculo automático) ──
    if doc_type == "CONSTANCIA_TERCIO":
        if _WP_TERCIO:
            logger.info(f"Proceso {process.id}: generando CONSTANCIA_TERCIO con WeasyPrint")
            inst    = _get_institution()
            student = _get_student(process.student_id)
            extra   = _get_extra_data(process)
            return generate_tercio_superior_weasyprint(process, student, extra, inst)
        else:
            logger.warning("WeasyPrint no instalado → usando ReportLab para CONSTANCIA_TERCIO")

    # ── CONSTANCIA_ORDEN_MERITO → HTML + WeasyPrint (cálculo automático) ──
    if doc_type == "CONSTANCIA_ORDEN_MERITO":
        if _WP_ORDEN:
            logger.info(f"Proceso {process.id}: generando CONSTANCIA_ORDEN_MERITO con WeasyPrint")
            inst    = _get_institution()
            student = _get_student(process.student_id)
            extra   = _get_extra_data(process)
            return generate_orden_merito_weasyprint(process, student, extra, inst)
        else:
            logger.warning("WeasyPrint no instalado → usando ReportLab para CONSTANCIA_ORDEN_MERITO")

    # ── FICHA_MATRICULA → HTML + WeasyPrint ──
    if doc_type == "FICHA_MATRICULA":
        if _WEASYPRINT_OK:
            logger.info(f"Proceso {process.id}: generando FICHA_MATRICULA con WeasyPrint")
            inst    = _get_institution()
            student = _get_student(process.student_id)
            extra   = _get_extra_data(process)
            courses = _get_enrolled_courses(student.get("plan_id"), _ciclo_int(student))
            return generate_ficha_matricula_weasyprint(process, student, extra, inst, courses)
        else:
            logger.warning("WeasyPrint no instalado → usando ReportLab para FICHA_MATRICULA")

    # ── CONSTANCIA_ESTUDIOS: intentar con plantilla .docx ──
    if doc_type == "CONSTANCIA_ESTUDIOS":
        template_ok = os.path.exists(CONSTANCIA_ESTUDIOS_TEMPLATE)
        lo_ok       = _find_libreoffice() is not None

        if HAS_DOCX and template_ok and lo_ok:
            logger.info(
                f"Proceso {process.id}: generando CONSTANCIA_ESTUDIOS "
                f"con plantilla .docx + LibreOffice"
            )
            inst    = _get_institution()
            student = _get_student(process.student_id)
            extra   = _get_extra_data(process)
            return _generate_constancia_docx(process, student, extra, inst)
        else:
            # Avisar en logs qué falta
            if not template_ok:
                logger.warning(
                    f"Plantilla .docx no encontrada en: {CONSTANCIA_ESTUDIOS_TEMPLATE} "
                    f"→ usando generador ReportLab"
                )
            elif not HAS_DOCX:
                logger.warning("python-docx no instalado → usando generador ReportLab")
            elif not lo_ok:
                logger.warning("LibreOffice no encontrado → usando generador ReportLab")

    # ── Resto de documentos (y fallback) → ReportLab ──
    if not HAS_REPORTLAB:
        raise ImportError("reportlab no está instalado. pip install reportlab")

    generator = DOCUMENT_GENERATORS.get(doc_type)
    if not generator:
        raise ValueError(
            f"No hay generador para '{doc_type}'. "
            f"Tipos soportados: {list(DOCUMENT_GENERATORS.keys())}"
        )

    inst    = _get_institution()
    student = _get_student(process.student_id)
    extra   = _get_extra_data(process)
    styles  = _get_styles()

    story  = generator(process, student, extra, styles, inst)
    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2 * cm,    bottomMargin=2 * cm,
    )
    doc.build(story)
    buffer.seek(0)

    now        = datetime.now()
    kind_clean = doc_type.replace("_", "-")
    filename   = f"{kind_clean}_{process.id:05d}_{now.strftime('%Y%m%d')}.pdf"
    return buffer, filename


# ═══════════════════════════════════════════════════════════════
# MAPEO Y LABELS
# ═══════════════════════════════════════════════════════════════

DOCUMENT_LABELS = {
    "CONSTANCIA_ESTUDIOS":     "Constancia de Estudios",
    "CONSTANCIA_ORDEN_MERITO": "Constancia de Orden de Mérito",
    "CONSTANCIA_TERCIO":       "Constancia de Tercio Superior",
    "CERTIFICADO_EGRESADO":    "Certificado de Egresado",
    "FICHA_MATRICULA":         "Ficha de Matrícula",
}


# ═══════════════════════════════════════════════════════════════
# VISTA API: GENERAR DOCUMENTO
# ═══════════════════════════════════════════════════════════════

class ProcessGenerateDocumentView(APIView):
    """
    POST /academic/processes/<pid>/generate-document

    Body (opcional):
      { "force": true }           → regenerar aunque ya exista
      { "document_type": "..." }  → tipo de documento
    """
    authentication_classes = [JWTAuthentication]
    permission_classes     = [permissions.IsAuthenticated]

    def post(self, request, pid):
        if not _can_admin_enroll(request.user):
            return Response(
                {"detail": "No tienes permisos para generar documentos"},
                status=status.HTTP_403_FORBIDDEN,
            )

        process = get_object_or_404(AcademicProcess, id=pid)
        body    = request.data or {}
        force   = body.get("force", False)

        doc_type = (body.get("document_type") or process.kind or "").upper().strip()

        if doc_type not in DOCUMENT_GENERATORS:
            return Response(
                {
                    "detail": f"Tipo de documento no soportado: '{doc_type}'",
                    "supported": list(DOCUMENT_GENERATORS.keys()),
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # ¿Ya existe?
        existing = ProcessFile.objects.filter(process_id=pid).order_by("-id")
        for f in existing:
            note = getattr(f, "note", "") or ""
            if "generado" in note.lower() or "auto" in note.lower():
                if not force:
                    return ok(
                        file=self._ser(f, request),
                        preview_url=self._abs(f, request),
                        already_exists=True,
                        message="Documento ya generado. Usa force=true para regenerar.",
                    )
                try:
                    if f.file:
                        f.file.delete(save=False)
                    f.delete()
                except Exception:
                    pass

        # Generar
        try:
            pdf_buffer, filename = generate_process_document(process, doc_type)
        except Exception as e:
            logger.error(f"Error generando documento proceso {pid}: {e}", exc_info=True)
            return Response(
                {"detail": f"Error al generar documento: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        label  = DOCUMENT_LABELS.get(doc_type, doc_type)
        kwargs = {
            "process_id": pid,
            "file": ContentFile(pdf_buffer.read(), name=filename),
        }
        if hasattr(ProcessFile, "note"):
            kwargs["note"] = f"Generado automáticamente — {label}"

        pf = ProcessFile.objects.create(**kwargs)

        return ok(
            file=self._ser(pf, request),
            preview_url=self._abs(pf, request),
            generated=True,
            document_type=doc_type,
            message=f"{label} generado exitosamente",
        )

    def _ser(self, f, request=None):
        url = f.file.url if f.file else ""
        return {
            "id":           f.id,
            "name":         (f.file.name or "").split("/")[-1] if f.file else "",
            "url":          url,
            "absolute_url": request.build_absolute_uri(url) if request and url else url,
            "size":         getattr(f.file, "size", 0),
            "note":         getattr(f, "note", "") or "",
        }

    def _abs(self, f, request=None):
        url = f.file.url if f.file else ""
        return request.build_absolute_uri(url) if request and url else url